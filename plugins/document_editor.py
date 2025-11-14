from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
from docx.shared import Inches
import os


D_ROUTE = "documents/"


class NewParagraphDeclined(Exception): pass

class WordDoc:
    def __init__(self, name: str | None = None):
        if not name:
            name = datetime.now().strftime("%d-%m-%Y")
        self.name = name
        
        try:
            self.document = Document(D_ROUTE + self.name + ".docx")
        except (FileNotFoundError, PackageNotFoundError):
            self.document = Document()

    def save(self):
        # логировать создание файла после импорта утилиты
        self.document.save(D_ROUTE + self.name + ".docx")

    def heading_exists(self, headline: str) -> bool:
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip() == headline.strip() and paragraph.style.name.startswith('Heading'):
                return True
        return False

    def new_paragraph(self, headline: str = "?", text: str = "..."):
        """returns false if the paragraph was declined"""
        if self.heading_exists(headline):
            raise NewParagraphDeclined()
        self.document.add_heading(headline, level=1)
        self.document.add_paragraph(text)
    
    def find_heading_index(self, headline: str) -> int:
        for i, paragraph in enumerate(self.document.paragraphs):
            if paragraph.text.strip() == headline.strip() and paragraph.style.name.startswith('Heading'):
                return i
        return -1
    
    def remove_paragraph(self, headline: str):
        heading_index = self.find_heading_index(headline)
        
        if heading_index == -1:
            return
        
        heading_element = self.document.paragraphs[heading_index]._element
        heading_element.getparent().remove(heading_element)
        
        if heading_index < len(self.document.paragraphs):
            text_element = self.document.paragraphs[heading_index]._element
            text_element.getparent().remove(text_element)
        
        self.save()

    def extract_all_content(self, filename: str | None = None):
        """
        Извлекает все содержимое документа в виде XML-строки
        """
        if filename:
            doc = Document(D_ROUTE + filename + ".docx")
        else:
            doc = self.document
            
        return doc.element.body.xml

    def _copy_paragraph(self, paragraph):
        new_paragraph = self.document.add_paragraph(style=paragraph.style)
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            # Копируем свойства форматирования
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

    def _copy_table(self, table):
        """Копирует таблицу с сохранением структуры и содержимого"""
        new_table = self.document.add_table(rows=len(table.rows), cols=len(table.columns))
        
        if table.style:
            new_table.style = table.style
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                new_cell.text = ""
                for paragraph in cell.paragraphs:
                    self._copy_paragraph_to_cell(paragraph, new_cell)

    def _copy_paragraph_to_cell(self, paragraph, cell):
        """Копирует параграф в ячейку таблицы"""
        new_paragraph = cell.add_paragraph(style=paragraph.style)
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

    def copy_all_content_from(self, donor_doc):
        """Копирует все содержимое из другого документа WordDoc"""
        donor = donor_doc.document
        
        for paragraph in donor.paragraphs:
            self._copy_paragraph(paragraph)
        
        for table in donor.tables:
            self._copy_table(table)
        
        if donor.sections and self.document.sections:
            self._copy_section_properties(donor.sections[0], self.document.sections[0])

    def _copy_section_properties(self, source_section, target_section):
        """Копирует свойства раздела"""
        target_section.page_height = source_section.page_height
        target_section.page_width = source_section.page_width
        target_section.left_margin = source_section.left_margin
        target_section.right_margin = source_section.right_margin
        target_section.top_margin = source_section.top_margin
        target_section.bottom_margin = source_section.bottom_margin
            
def extract_text(file: WordDoc):
    full_text = []
    for paragraph in file.document.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    return '\n'.join(full_text)

def merge_files(donor: WordDoc, target: WordDoc):
    """
    Улучшенная функция слияния файлов, которая правильно копирует таблицы
    """
    target.copy_all_content_from(donor)
    target.save()
    print(f"Все содержимое из '{donor.name}.docx' скопировано в '{target.name}.docx'")
    

if __name__ == "__main__":
    try:
        t = WordDoc("conclusion")
        d = WordDoc("lavrova")

        merge_files(d, t)
        
    except Exception as e:
        print(e)
